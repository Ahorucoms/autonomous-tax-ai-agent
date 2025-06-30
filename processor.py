"""
Document Processing Service for Malta Tax AI
Handles document parsing, OCR, data extraction, and integration with vector storage
"""

import os
import io
import json
import logging
import asyncio
from typing import Dict, List, Any, Optional, Tuple, Union
from datetime import datetime
import hashlib
import mimetypes
from pathlib import Path

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import re
from dataclasses import dataclass
from enum import Enum


class DocumentType(Enum):
    """Supported document types"""
    FS3 = "fs3"  # Individual income tax return
    FS5 = "fs5"  # Corporate tax return
    VAT_RETURN = "vat_return"
    PAYSLIP = "payslip"
    BANK_STATEMENT = "bank_statement"
    INVOICE = "invoice"
    RECEIPT = "receipt"
    CONTRACT = "contract"
    REGULATION = "regulation"
    GUIDANCE = "guidance"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """Document processing status"""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    PARTIAL = "partial"


@dataclass
class DocumentMetadata:
    """Document metadata structure"""
    filename: str
    file_type: str
    file_size: int
    document_type: DocumentType
    language: str = "en"
    pages: int = 0
    confidence: float = 0.0
    processing_time: float = 0.0
    created_at: datetime = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.utcnow()


@dataclass
class ExtractedData:
    """Structured data extracted from documents"""
    fields: Dict[str, Any]
    tables: List[Dict[str, Any]]
    entities: List[Dict[str, Any]]
    confidence_scores: Dict[str, float]
    validation_errors: List[str]


@dataclass
class ProcessingResult:
    """Complete document processing result"""
    document_id: str
    metadata: DocumentMetadata
    raw_text: str
    structured_text: str
    extracted_data: ExtractedData
    status: ProcessingStatus
    error_message: Optional[str] = None


class DocumentProcessor:
    """Main document processing service"""
    
    def __init__(self, 
                 temp_dir: str = "/tmp/document_processing",
                 max_file_size: int = 50 * 1024 * 1024):  # 50MB
        """
        Initialize document processor
        
        Args:
            temp_dir: Temporary directory for processing
            max_file_size: Maximum file size in bytes
        """
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(exist_ok=True)
        self.max_file_size = max_file_size
        self.logger = logging.getLogger(__name__)
        
        # Initialize extractors
        self.extractors = {
            DocumentType.FS3: FS3Extractor(),
            DocumentType.FS5: FS5Extractor(),
            DocumentType.VAT_RETURN: VATReturnExtractor(),
            DocumentType.PAYSLIP: PayslipExtractor(),
            DocumentType.INVOICE: InvoiceExtractor(),
        }
    
    async def process_document(self, 
                             file_data: bytes,
                             filename: str,
                             document_type: Optional[DocumentType] = None) -> ProcessingResult:
        """
        Process a document from bytes
        
        Args:
            file_data: Document file data
            filename: Original filename
            document_type: Detected or specified document type
            
        Returns:
            Processing result
        """
        start_time = datetime.utcnow()
        document_id = self._generate_document_id(filename, file_data)
        
        try:
            # Validate file size
            if len(file_data) > self.max_file_size:
                raise ValueError(f"File size {len(file_data)} exceeds maximum {self.max_file_size}")
            
            # Detect file type
            file_type = self._detect_file_type(filename, file_data)
            
            # Auto-detect document type if not provided
            if document_type is None:
                document_type = self._detect_document_type(filename, file_data)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=filename,
                file_type=file_type,
                file_size=len(file_data),
                document_type=document_type
            )
            
            # Extract text content
            raw_text, structured_text = await self._extract_text(file_data, file_type)
            
            # Extract structured data
            extracted_data = await self._extract_structured_data(
                raw_text, structured_text, document_type
            )
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            metadata.processing_time = processing_time
            
            # Determine overall confidence
            if extracted_data.confidence_scores:
                metadata.confidence = sum(extracted_data.confidence_scores.values()) / len(extracted_data.confidence_scores)
            
            # Create result
            result = ProcessingResult(
                document_id=document_id,
                metadata=metadata,
                raw_text=raw_text,
                structured_text=structured_text,
                extracted_data=extracted_data,
                status=ProcessingStatus.COMPLETED
            )
            
            self.logger.info(f"Document {document_id} processed successfully in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            self.logger.error(f"Document processing failed: {e}")
            
            # Return failed result
            return ProcessingResult(
                document_id=document_id,
                metadata=DocumentMetadata(
                    filename=filename,
                    file_type="unknown",
                    file_size=len(file_data),
                    document_type=DocumentType.UNKNOWN
                ),
                raw_text="",
                structured_text="",
                extracted_data=ExtractedData(
                    fields={},
                    tables=[],
                    entities=[],
                    confidence_scores={},
                    validation_errors=[str(e)]
                ),
                status=ProcessingStatus.FAILED,
                error_message=str(e)
            )
    
    def _generate_document_id(self, filename: str, file_data: bytes) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.sha256(file_data).hexdigest()[:16]
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        return f"doc_{timestamp}_{content_hash}"
    
    def _detect_file_type(self, filename: str, file_data: bytes) -> str:
        """Detect file MIME type"""
        # Try to detect from filename
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            return mime_type
        
        # Try to detect from file signature
        if file_data.startswith(b'%PDF'):
            return 'application/pdf'
        elif file_data.startswith(b'PK'):
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_data.startswith(b'\xff\xd8\xff'):
            return 'image/jpeg'
        elif file_data.startswith(b'\x89PNG'):
            return 'image/png'
        
        return 'application/octet-stream'
    
    def _detect_document_type(self, filename: str, file_data: bytes) -> DocumentType:
        """Auto-detect document type from filename and content"""
        filename_lower = filename.lower()
        
        # Check filename patterns
        if 'fs3' in filename_lower or 'income_tax' in filename_lower:
            return DocumentType.FS3
        elif 'fs5' in filename_lower or 'corporate_tax' in filename_lower:
            return DocumentType.FS5
        elif 'vat' in filename_lower:
            return DocumentType.VAT_RETURN
        elif 'payslip' in filename_lower or 'salary' in filename_lower:
            return DocumentType.PAYSLIP
        elif 'invoice' in filename_lower:
            return DocumentType.INVOICE
        elif 'receipt' in filename_lower:
            return DocumentType.RECEIPT
        
        return DocumentType.UNKNOWN
    
    async def _extract_text(self, file_data: bytes, file_type: str) -> Tuple[str, str]:
        """Extract text from document"""
        try:
            if file_type == 'application/pdf':
                return await self._extract_pdf_text(file_data)
            elif file_type.startswith('image/'):
                return await self._extract_image_text(file_data)
            elif 'wordprocessingml' in file_type:
                return await self._extract_docx_text(file_data)
            elif file_type.startswith('text/'):
                text = file_data.decode('utf-8', errors='ignore')
                return text, text
            else:
                # Try to decode as text
                text = file_data.decode('utf-8', errors='ignore')
                return text, text
                
        except Exception as e:
            self.logger.error(f"Text extraction failed: {e}")
            return "", ""
    
    async def _extract_pdf_text(self, file_data: bytes) -> Tuple[str, str]:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            return "[PDF parsing not available]", "[PDF parsing not available]"
        
        try:
            # Use pdfplumber for better text extraction
            with io.BytesIO(file_data) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    raw_text = ""
                    structured_text = ""
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        raw_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            structured_text += f"\n--- Page {page_num + 1} Tables ---\n"
                            for table_num, table in enumerate(tables):
                                structured_text += f"Table {table_num + 1}:\n"
                                for row in table:
                                    structured_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                        
                        structured_text += page_text + "\n"
            
            return raw_text.strip(), structured_text.strip()
            
        except Exception as e:
            self.logger.error(f"PDF text extraction failed: {e}")
            # Fallback to PyPDF2
            try:
                with io.BytesIO(file_data) as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip(), text.strip()
            except:
                return "[PDF text extraction failed]", "[PDF text extraction failed]"
    
    async def _extract_image_text(self, file_data: bytes) -> Tuple[str, str]:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            return "[OCR not available]", "[OCR not available]"
        
        try:
            # Open image
            image = Image.open(io.BytesIO(file_data))
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            return text.strip(), text.strip()
            
        except Exception as e:
            self.logger.error(f"OCR extraction failed: {e}")
            return "[OCR extraction failed]", "[OCR extraction failed]"
    
    async def _extract_docx_text(self, file_data: bytes) -> Tuple[str, str]:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            return "[DOCX parsing not available]", "[DOCX parsing not available]"
        
        try:
            doc = DocxDocument(io.BytesIO(file_data))
            
            raw_text = ""
            structured_text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    raw_text += text + "\n"
                    structured_text += text + "\n"
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                structured_text += f"\n--- Table {table_num + 1} ---\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    structured_text += row_text + "\n"
                    raw_text += row_text + "\n"
            
            return raw_text.strip(), structured_text.strip()
            
        except Exception as e:
            self.logger.error(f"DOCX text extraction failed: {e}")
            return "[DOCX extraction failed]", "[DOCX extraction failed]"
    
    async def _extract_structured_data(self, 
                                     raw_text: str,
                                     structured_text: str,
                                     document_type: DocumentType) -> ExtractedData:
        """Extract structured data based on document type"""
        try:
            if document_type in self.extractors:
                extractor = self.extractors[document_type]
                return await extractor.extract(raw_text, structured_text)
            else:
                # Generic extraction
                return await self._generic_extraction(raw_text, structured_text)
                
        except Exception as e:
            self.logger.error(f"Structured data extraction failed: {e}")
            return ExtractedData(
                fields={},
                tables=[],
                entities=[],
                confidence_scores={},
                validation_errors=[str(e)]
            )
    
    async def _generic_extraction(self, raw_text: str, structured_text: str) -> ExtractedData:
        """Generic data extraction for unknown document types"""
        fields = {}
        entities = []
        confidence_scores = {}
        
        # Extract common patterns
        
        # Dates
        date_pattern = r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b'
        dates = re.findall(date_pattern, raw_text)
        if dates:
            fields['dates'] = dates
            confidence_scores['dates'] = 0.8
        
        # Amounts (EUR)
        amount_pattern = r'€\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        amounts = re.findall(amount_pattern, raw_text)
        if amounts:
            fields['amounts'] = amounts
            confidence_scores['amounts'] = 0.9
        
        # Email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, raw_text)
        if emails:
            fields['emails'] = emails
            confidence_scores['emails'] = 0.95
        
        # Phone numbers
        phone_pattern = r'\b(?:\+356\s?)?(?:\d{4}\s?\d{4}|\d{8})\b'
        phones = re.findall(phone_pattern, raw_text)
        if phones:
            fields['phone_numbers'] = phones
            confidence_scores['phone_numbers'] = 0.7
        
        return ExtractedData(
            fields=fields,
            tables=[],
            entities=entities,
            confidence_scores=confidence_scores,
            validation_errors=[]
        )


class BaseExtractor:
    """Base class for document-specific extractors"""
    
    async def extract(self, raw_text: str, structured_text: str) -> ExtractedData:
        """Extract structured data from text"""
        raise NotImplementedError


class FS3Extractor(BaseExtractor):
    """Extractor for FS3 (Individual Income Tax Return) forms"""
    
    async def extract(self, raw_text: str, structured_text: str) -> ExtractedData:
        fields = {}
        confidence_scores = {}
        validation_errors = []
        
        # Extract taxpayer name
        name_pattern = r'(?:name|taxpayer)[:\s]+([A-Za-z\s]+)'
        name_match = re.search(name_pattern, raw_text, re.IGNORECASE)
        if name_match:
            fields['taxpayer_name'] = name_match.group(1).strip()
            confidence_scores['taxpayer_name'] = 0.85
        
        # Extract ID number
        id_pattern = r'(?:id|identity)[:\s]+(\d+[A-Z]?)'
        id_match = re.search(id_pattern, raw_text, re.IGNORECASE)
        if id_match:
            fields['id_number'] = id_match.group(1).strip()
            confidence_scores['id_number'] = 0.9
        
        # Extract annual income
        income_pattern = r'(?:annual\s+income|total\s+income)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        income_match = re.search(income_pattern, raw_text, re.IGNORECASE)
        if income_match:
            fields['annual_income'] = income_match.group(1).replace(',', '')
            confidence_scores['annual_income'] = 0.9
        
        # Extract tax year
        year_pattern = r'(?:tax\s+year|year)[:\s]+(\d{4})'
        year_match = re.search(year_pattern, raw_text, re.IGNORECASE)
        if year_match:
            fields['tax_year'] = int(year_match.group(1))
            confidence_scores['tax_year'] = 0.95
        
        # Extract marital status
        marital_pattern = r'(?:marital\s+status|status)[:\s]+(single|married|divorced|widowed)'
        marital_match = re.search(marital_pattern, raw_text, re.IGNORECASE)
        if marital_match:
            fields['marital_status'] = marital_match.group(1).lower()
            confidence_scores['marital_status'] = 0.8
        
        # Validation
        if 'annual_income' in fields:
            try:
                income = float(fields['annual_income'])
                if income < 0:
                    validation_errors.append("Annual income cannot be negative")
                elif income > 10000000:  # 10M EUR seems unreasonable
                    validation_errors.append("Annual income seems unreasonably high")
            except ValueError:
                validation_errors.append("Invalid annual income format")
        
        if 'tax_year' in fields:
            year = fields['tax_year']
            current_year = datetime.now().year
            if year < 2000 or year > current_year + 1:
                validation_errors.append(f"Invalid tax year: {year}")
        
        return ExtractedData(
            fields=fields,
            tables=[],
            entities=[],
            confidence_scores=confidence_scores,
            validation_errors=validation_errors
        )


class FS5Extractor(BaseExtractor):
    """Extractor for FS5 (Corporate Tax Return) forms"""
    
    async def extract(self, raw_text: str, structured_text: str) -> ExtractedData:
        fields = {}
        confidence_scores = {}
        validation_errors = []
        
        # Extract company name
        company_pattern = r'(?:company\s+name|entity)[:\s]+([A-Za-z\s&.,]+)'
        company_match = re.search(company_pattern, raw_text, re.IGNORECASE)
        if company_match:
            fields['company_name'] = company_match.group(1).strip()
            confidence_scores['company_name'] = 0.85
        
        # Extract registration number
        reg_pattern = r'(?:registration|reg\.?\s+no\.?)[:\s]+([A-Z]?\d+)'
        reg_match = re.search(reg_pattern, raw_text, re.IGNORECASE)
        if reg_match:
            fields['registration_number'] = reg_match.group(1).strip()
            confidence_scores['registration_number'] = 0.9
        
        # Extract turnover
        turnover_pattern = r'(?:turnover|revenue|sales)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        turnover_match = re.search(turnover_pattern, raw_text, re.IGNORECASE)
        if turnover_match:
            fields['turnover'] = turnover_match.group(1).replace(',', '')
            confidence_scores['turnover'] = 0.85
        
        # Extract profit
        profit_pattern = r'(?:profit|net\s+income)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        profit_match = re.search(profit_pattern, raw_text, re.IGNORECASE)
        if profit_match:
            fields['profit'] = profit_match.group(1).replace(',', '')
            confidence_scores['profit'] = 0.85
        
        return ExtractedData(
            fields=fields,
            tables=[],
            entities=[],
            confidence_scores=confidence_scores,
            validation_errors=validation_errors
        )


class VATReturnExtractor(BaseExtractor):
    """Extractor for VAT Return forms"""
    
    async def extract(self, raw_text: str, structured_text: str) -> ExtractedData:
        fields = {}
        confidence_scores = {}
        validation_errors = []
        
        # Extract VAT number
        vat_pattern = r'(?:vat\s+no\.?|vat\s+number)[:\s]+(\d+)'
        vat_match = re.search(vat_pattern, raw_text, re.IGNORECASE)
        if vat_match:
            fields['vat_number'] = vat_match.group(1).strip()
            confidence_scores['vat_number'] = 0.95
        
        # Extract period
        period_pattern = r'(?:period|quarter)[:\s]+([Q1-4\s\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{1,2}/\d{4})'
        period_match = re.search(period_pattern, raw_text, re.IGNORECASE)
        if period_match:
            fields['period'] = period_match.group(1).strip()
            confidence_scores['period'] = 0.8
        
        # Extract total VAT
        total_vat_pattern = r'(?:total\s+vat|vat\s+due)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        total_vat_match = re.search(total_vat_pattern, raw_text, re.IGNORECASE)
        if total_vat_match:
            fields['total_vat'] = total_vat_match.group(1).replace(',', '')
            confidence_scores['total_vat'] = 0.9
        
        return ExtractedData(
            fields=fields,
            tables=[],
            entities=[],
            confidence_scores=confidence_scores,
            validation_errors=validation_errors
        )


class PayslipExtractor(BaseExtractor):
    """Extractor for payslip documents"""
    
    async def extract(self, raw_text: str, structured_text: str) -> ExtractedData:
        fields = {}
        confidence_scores = {}
        
        # Extract employee name
        name_pattern = r'(?:employee|name)[:\s]+([A-Za-z\s]+)'
        name_match = re.search(name_pattern, raw_text, re.IGNORECASE)
        if name_match:
            fields['employee_name'] = name_match.group(1).strip()
            confidence_scores['employee_name'] = 0.8
        
        # Extract gross salary
        gross_pattern = r'(?:gross\s+salary|gross\s+pay)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        gross_match = re.search(gross_pattern, raw_text, re.IGNORECASE)
        if gross_match:
            fields['gross_salary'] = gross_match.group(1).replace(',', '')
            confidence_scores['gross_salary'] = 0.9
        
        # Extract net salary
        net_pattern = r'(?:net\s+salary|net\s+pay|take\s+home)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        net_match = re.search(net_pattern, raw_text, re.IGNORECASE)
        if net_match:
            fields['net_salary'] = net_match.group(1).replace(',', '')
            confidence_scores['net_salary'] = 0.9
        
        return ExtractedData(
            fields=fields,
            tables=[],
            entities=[],
            confidence_scores=confidence_scores,
            validation_errors=[]
        )


class InvoiceExtractor(BaseExtractor):
    """Extractor for invoice documents"""
    
    async def extract(self, raw_text: str, structured_text: str) -> ExtractedData:
        fields = {}
        confidence_scores = {}
        
        # Extract invoice number
        invoice_pattern = r'(?:invoice\s+no\.?|invoice\s+number)[:\s]+([A-Z0-9-]+)'
        invoice_match = re.search(invoice_pattern, raw_text, re.IGNORECASE)
        if invoice_match:
            fields['invoice_number'] = invoice_match.group(1).strip()
            confidence_scores['invoice_number'] = 0.9
        
        # Extract total amount
        total_pattern = r'(?:total|amount\s+due)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        total_match = re.search(total_pattern, raw_text, re.IGNORECASE)
        if total_match:
            fields['total_amount'] = total_match.group(1).replace(',', '')
            confidence_scores['total_amount'] = 0.85
        
        # Extract VAT amount
        vat_pattern = r'(?:vat|tax)[:\s]+€?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
        vat_match = re.search(vat_pattern, raw_text, re.IGNORECASE)
        if vat_match:
            fields['vat_amount'] = vat_match.group(1).replace(',', '')
            confidence_scores['vat_amount'] = 0.8
        
        return ExtractedData(
            fields=fields,
            tables=[],
            entities=[],
            confidence_scores=confidence_scores,
            validation_errors=[]
        )


# Example usage and testing
if __name__ == "__main__":
    import asyncio
    
    async def test_processor():
        # Initialize processor
        processor = DocumentProcessor()
        
        # Test with sample text document
        sample_text = """
        FS3 - Individual Income Tax Return
        
        Taxpayer Name: John Doe
        ID Number: 123456M
        Annual Income: €45,000.00
        Tax Year: 2025
        Marital Status: Single
        
        This is a sample FS3 form for testing purposes.
        """
        
        # Process document
        result = await processor.process_document(
            file_data=sample_text.encode('utf-8'),
            filename="sample_fs3.txt",
            document_type=DocumentType.FS3
        )
        
        print(f"Processing Status: {result.status}")
        print(f"Document ID: {result.document_id}")
        print(f"Extracted Fields: {result.extracted_data.fields}")
        print(f"Confidence Scores: {result.extracted_data.confidence_scores}")
        print(f"Validation Errors: {result.extracted_data.validation_errors}")
    
    # Run test
    asyncio.run(test_processor())

